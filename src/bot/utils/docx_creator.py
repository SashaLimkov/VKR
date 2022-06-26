import os
from time import sleep

from docx import Document
from docx2pdf import convert

from usersupport.models import Client

import secrets
import string


def generate_alphanum_crypt_string():
    letters_and_digits = string.ascii_letters + string.digits
    crypt_rand_string = ''.join(secrets.choice(
        letters_and_digits) for i in range(5))
    return crypt_rand_string


def get_docx(user: Client):
    document = Document()
    print(user.prediction)
    document.add_paragraph(f"{user.user.name}")
    document.add_paragraph("")
    document.add_paragraph(f"'Контактные данные:'")
    p = document.add_paragraph(
        f"Номер телефона: {user.user.phone}'", style='List Number'
    )
    p = document.add_paragraph(
        f'Почта: {user.user.email}', style='List Number'
    )
    p = document.add_paragraph(
        f'Антропометрические данные: '
    )
    p = document.add_paragraph(
        f'Возраст: {user.age}', style='List Bullet'
    )
    p = document.add_paragraph(
        f'Рост: {user.height}', style='List Bullet'
    )
    p = document.add_paragraph(
        f'Вес: {user.weight}', style='List Bullet'
    )
    p = document.add_paragraph(
        f'ИМС: {user.ims}', style='List Bullet'
    )
    p = document.add_paragraph("")
    p = document.add_paragraph(f"Жалобы пользователя: {user.additional if user.additional else ''}")
    p = document.add_paragraph(f"Предварительный диагноз и рекомендации: {user.prediction if user.prediction else ''}")
    document.save(f'{user.user.name}.docx')
    key = generate_alphanum_crypt_string()
    convert(f'{user.user.name}.docx', f'{user.user.name} {key}.pdf')
    sleep(5)
    os.remove(f'{user.user.name}.docx')
    # await asyncio.sleep(20)
    # file.close()
    return f"{user.user.name} {key}.pdf"
